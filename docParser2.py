# -*- coding: utf-8 -*-
"""
Created on Sat Sep  7 20:32:03 2019

@author: lenovo
"""


import os.path
#import win32com.client
import docx
import os
import re
import sqlite3
import MeCab
import calendar
import subprocess
import glob
import logging

from sql.connect import DBConnection

logger = logging.getLogger('doc')
logger.setLevel(logging.INFO)
formatter = logging.Formatter('%(levelname)s:%(message)s:%(asctime)s')
file_handler = logging.FileHandler('docParser.log')
file_handler.setFormatter(formatter)
logger.addHandler(file_handler)

main_logger = logging.getLogger('main')
main_logger.setLevel(logging.INFO)
main_logger_format = logging.Formatter('%(message)s:%(asctime)s')
main_logger_handler = logging.FileHandler('MainLog.log')
main_logger_handler.setFormatter(main_logger_format)
main_logger.addHandler(main_logger_handler)

db = DBConnection()
conn = db.connection()
c = conn.cursor()
mecab = MeCab.Tagger()

#baseDir = '/home/ihsan/Nikkei/news/20130408朝刊' 
baseDir = os.path.abspath(os.path.dirname(__file__))

#News
art_id = []
art_ver = []
art_titles = []
art_subtitles = []
art_content = []
art_day = []
art_year = []
art_month = []
art_date = []
art_date_full = []
art_raw = []
art_category_id = []

#Media
art_media = []
art_media_type = 'Picture'

#Categories
#morning_categories = []
#morning_categories_id = []
#evening_categories = []
#evening_categories_id = []
#new_category = []
#cat_ver = []
#
##Days
#day_id = [1,2,3,4,5,6,7]
#day_kanji = ['日','月','火','水','木','金','土']

#Dictionary
artikel = []
title_cleaned = []
subtitle_cleaned = []
content_cleaned = []

#Parser
japanese = []
furigana = []
word_type = []
row=dict()

directories = []

def doc2docx(dir):
    os.chdir(dir)

    for doc in glob.iglob("*.doc"):

        #print(doc)
        subprocess.call(['soffice', '--headless', '--convert-to', 'docx', doc], shell = False)

    
###################################################################################################
    
def cleaner(container, string):
    words = []
    lines = mecab.parse(string).split('\n')
    for line in lines:
        items = re.split('[\t,]',line)
        words.append(items[0])
    a = ' '.join(words)    
    container.append(a[:-5])    
    
def get_dir(baseDir):
    doc2docx(baseDir)
    for entry in os.listdir(baseDir):
        path = os.path.join(baseDir, entry)
        if os.path.isdir(path):
            get_dir(path)
        else:
            #print(entry)
            if entry[-5:] == '.docx':
                if entry[0:4] != '~$13': 
                    directories.append(path)
    #return directories
    
    
def docParser(dirs):
    for j in range(len(dirs)):
        doc = docx.Document(dirs[j])
    #TITLE
        header = doc.paragraphs[0].text
        start = header.find('刊')
        header_len = len(header)
        temp_title = header[start+1:header_len].split(' ')
        title = temp_title[0]
        print(title)
        cleaner(title_cleaned, title)
    #SUBTITLE
        try:
            subtitle = temp_title[1]
        except:
            subtitle = None
        cleaner(subtitle_cleaned, str(subtitle))
    #TEXT CONTENT    
        temp2 = []
        cleaned_temp = []
        length = len(doc.paragraphs)
        content = doc.paragraphs[3:length]
        for i in content:
            temp2.append(i.text)
            cleaner(cleaned_temp, i.text)
            
        content_cleaned.append(cleaned_temp)
            
    #DATE FULL    
        header = doc.paragraphs[0].text
        year = header[:4]
        getsu = header.find('月')
        nen = header.find('年')
        hi = header.find('日')
        month = f'{header[nen+1:getsu]:0>2}'
        date = f'{header[getsu+1:hi]:0>2}'
        date_full = f'{year}-{month}-{date}'
        day_id = calendar.weekday(int(year), int(month), int(date))
        if day_id!=6:
            day_id = day_id+2
        else:
            day_id = 1
        
        
    #VERSION    
        publication = header.find('刊')
        ver = header[hi+1:publication+1] 
        if ver == '朝刊':
            version = 'M'
        elif ver == '夕刊':
            version = 'E'
    
#ID        
        ids = f'{year}{month}{date}{version}{j+1:0>4}'
        
        
        if ids:
            art_id.append(ids)
        if version:
            art_ver.append(version)
        if title:    
            art_titles.append(title)
        if temp2:
            art_content.append(''.join(temp2))            
        art_subtitles.append(subtitle)
        if year:
            art_year.append(year)
        if month:
            art_month.append(month)
        if date:
            art_date.append(date)
        if date_full:
            art_date_full.append(date_full)
        if day_id:
            art_day.append(day_id)
        

        logger.info(f'Article fetched: date = {date_full}, version = {version}, id = {ids}')
###########################################################################################################
    
    for i in range(len(content_cleaned)):
        artikel.append([title_cleaned[i], subtitle_cleaned[i], ''.join(content_cleaned[i])])    
    
    for i in range(len(artikel)):
        artikel[i] = ' '.join(artikel[i])
        
    
###########################################################################################################    
    for i in range(len(artikel)):
        node = mecab.parse(artikel[i])
        node = node.split("\n")
        for j in node:
            feature = re.split('[\t,]',j)
            if len(feature) >= 2 and feature[1] in ('助詞', '記号'):
                continue
            if len(feature) >= 2 and feature[2] in ('数'):
                continue
            if len(feature) <= 8:
                continue
            if feature[0] == "EOS":
                break
            japanese.append(feature[0])
            furigana.append(feature[8])
            word_type.append(feature[1])
    
######################################################################################################    
    x=[]
    for i in range(len(artikel)):
        one = mecab.parse(artikel[i])
        x.append(one)
    temp = []
    for i in x:
        a=i.split('\n')
        temp.append(a)
    temp2=[]
    for i in temp:
        f=[]
        for j in i:
            if len(j)>1:
                b= re.split('[\t,]',j)
                if len(b) >= 2 and b[1] in ('助詞', '記号'):
                    continue
                if len(b) >= 2 and b[2] in ('数'):
                    continue
                if len(b) <= 8:
                    continue
                if b[0] == 'EOS':
                    break
                f.append(b[0])
        temp2.append(f)
        
    def word_count(str):
        counts = dict()    
        for japanese in str:
            if japanese in counts:
                counts[japanese] += 1
            else:
                counts[japanese] = 1
        return counts

    temp3=[]
    for i in temp2:
        o = word_count(i)
        temp3.append(o)
        

    for i in range(len(art_id)):
        row[art_id[i]] = temp3[i]    
 
######################################################################################################    

       
def main(baseDir):
    #doc2docx(baseDir)
    #dirs = get_dir(baseDir)
    #docParser(dirs)
    get_dir(baseDir)
    docParser(directories)
    
    #print(directories)
    
    print('Inserting to News, Media ....')
    logger.info('Inserting to News, Media ....')
    for i in range(len(art_titles)):
        c.execute("insert or ignore into news (News_ID, Date_Full, Year, Month, Date, Day, Version, Title, Subtitle, Text_Content) values(?, ?, ?, ?, ?, ?, ?, ?, ?, ?)", (art_id[i], art_date_full[i], int(art_year[i]), int(art_month[i]), int(art_date[i]), int(art_day[i]), art_ver[i], art_titles[i], art_subtitles[i], str(art_content[i])))
        c.execute("insert or ignore into media (News_ID, type) values (?,?)", (art_id[i], art_media_type))
        conn.commit()
    logger.info('Finished Inserting to News, Media ....')
    print('Finished Inserting to News, Media')    

    print('Inserting to Dictionary ....')
    logger.info('Inserting to Dictionary ....')
    for i in range(len(japanese)):
        c.execute("insert or ignore into dictionary (japanese, furigana, word_type) values (?,?,?)", (japanese[i], furigana[i], word_type[i]))
        conn.commit()
    print('Finished Inserting to Dictionary')
    logger.info('Finished Inserting to Dictionary')

    a = c.execute('select word_id, japanese from dictionary')    

    name = []
    ids = []
    for i in a:
        ids.append(i[0])   
        name.append(i[1])
    

    print('Inserting into word_count ...')
    logger.info('Inserting into word_count ...')
    for key, value in row.items():
        a = key
        b = value
        for key, value in b.items():
            for i in range(len(name)):
                if key == name[i]:
                    key = ids[i]
            c.execute('insert or ignore into word_count(word_id, news_id, count) values (?,?,?)', (key, a, value))
            conn.commit()
    print('Finished Inserting into word_count')
    logger.info('Finished Inserting into word_count')


    b = c.execute('select word_id, sum(count) from word_count group by word_id')
    ids = []
    count = []
    for i in b:
        ids.append(i[0])
        count.append(i[1])
        
    print('Updating Count Total ....')
    logger.info('Updating Count Total ....')
    for i in range(len(ids)):
        c.execute('UPDATE dictionary SET count_total = %d where word_id = %s' %(count[i], ids[i]))
        conn.commit()
    print('Finished Updating Count Total')
    logger.info('Finished Updating Count Total')
    conn.close()
    
    main_logger.warning(f'Articles from {art_date_full[0]} version {art_ver[0]} has been added to database @ ')

    print('removing duplicates')    
    for i in directories:
        os.remove(i)
        
        
######################################################################################################                        
            
main(baseDir)
    
    
    
    
    










